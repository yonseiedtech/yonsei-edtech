# -*- coding: utf-8 -*-
"""
theory-implementation-matrix.draft.md → .docx 변환 스크립트.
python-docx 1.2.0 기반 line-by-line markdown 파서.

지원 요소:
- # / ## / ### → Heading 1/2/3
- > ![alt](path) → image insert + caption paragraph (italic, centered)
- ![alt](path) → image insert (centered)
- > (그 외) → indented paragraph
- - / * → bullet
- 1. / 2. → numbered list
- ``` ... ``` → monospace code block
- --- → horizontal rule (thematic break)
- **bold** → bold inline
- *italic* → italic inline
- `code` → monospace inline
- (그 외) → 일반 단락
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "theory-implementation-matrix.draft.md")
DST = os.path.join(HERE, "theory-implementation-matrix.draft.docx")


def set_default_font(doc, font_name="Malgun Gothic", size=11):
    """기본 폰트 — 한글 친화"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    # 한글 폰트 명시 (rFonts eastAsia)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_inline_runs(paragraph, text):
    """단순 inline 형식 처리: **bold**, *italic*, `code`, [text](url)"""
    # 매우 단순한 토크나이저
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and not token.startswith("**"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("["):
            # [label](url) — 단순 hyperlink 텍스트로 표시
            mlink = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mlink:
                run = paragraph.add_run(mlink.group(1))
                run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
                run.underline = True
            else:
                paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_image_paragraph(doc, alt, path):
    """이미지 삽입 — 가운데 정렬 + 캡션 분리 처리"""
    full = os.path.normpath(os.path.join(HERE, path))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        if os.path.exists(full):
            run.add_picture(full, width=Cm(15))
        else:
            run.add_text(f"[Image not found: {path}]")
            run.italic = True
    except Exception as e:
        run.add_text(f"[Image error: {path} — {e}]")
        run.italic = True


def parse_markdown_to_docx(src_path, dst_path):
    """간단 파서. 한 라인씩 처리."""
    with open(src_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    set_default_font(doc)

    # 페이지 여백
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    in_code = False
    code_buf = []
    in_blockquote_caption = False  # > ![Figure ...] 다음의 caption 처리용
    last_image_paragraph = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # code fence
        if stripped.startswith("```"):
            if in_code:
                # close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # blockquote with image (figure)
        m_img_in_quote = re.match(r"^>\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m_img_in_quote:
            add_image_paragraph(doc, m_img_in_quote.group(1), m_img_in_quote.group(2))
            in_blockquote_caption = True
            i += 1
            continue

        # blockquote with text — caption or note
        if stripped.startswith(">"):
            quote_text = stripped[1:].lstrip()
            if not quote_text:
                in_blockquote_caption = False
                i += 1
                continue
            p = doc.add_paragraph()
            if in_blockquote_caption:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(quote_text)
                run.italic = True
                run.font.size = Pt(10)
            else:
                p.paragraph_format.left_indent = Cm(1)
                add_inline_runs(p, quote_text)
                for run in p.runs:
                    run.italic = True
            i += 1
            continue
        else:
            in_blockquote_caption = False

        # standalone image
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m_img:
            add_image_paragraph(doc, m_img.group(1), m_img.group(2))
            i += 1
            continue

        # horizontal rule
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 50).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # headings
        m_h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m_h:
            level = len(m_h.group(1))
            text = m_h.group(2).strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # bullet
        m_bul = re.match(r"^[\-\*]\s+(.*)$", stripped)
        if m_bul:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m_bul.group(1))
            i += 1
            continue

        # numbered list
        m_num = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m_num.group(1))
            i += 1
            continue

        # table — 시작 라인이 | 로 시작
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1
            # parse rows
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            # 두번째 행이 구분자(---)면 제거
            if len(rows) >= 2 and all(re.match(r"^[\s\-:]+$", c) for c in rows[1]):
                rows = [rows[0]] + rows[2:]
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Light Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci >= cols:
                            continue
                        cell_obj = table.rows[ri].cells[ci]
                        cell_obj.text = ""
                        p = cell_obj.paragraphs[0]
                        add_inline_runs(p, cell)
                        if ri == 0:
                            for run in p.runs:
                                run.bold = True
            i = j
            continue

        # empty line → paragraph break (이미 자동)
        if not stripped:
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(dst_path)
    print(f"saved: {dst_path}")


if __name__ == "__main__":
    parse_markdown_to_docx(SRC, DST)
